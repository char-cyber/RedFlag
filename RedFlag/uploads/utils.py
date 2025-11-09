from io import BytesIO
from pypdf import PdfReader
import fitz  # PyMuPDF
from docx import Document

def preprocess(file_obj):
    """
    Analyze file (PDF or DOCX):
    - Check validity
    - Count pages or sections
    - Count images
    - Extract basic metadata
    """

    result = {
        "is_pdf": False,
        "is_docx": False,
        "num_pages": 0,
        "num_images": 0,
        "errors": [],
        "file_type": None
    }

    filename = getattr(file_obj, "name", "").lower()

    try:
        # Read file into memory once (works for Django InMemoryUploadedFile)
        file_data = file_obj.read()
        buffer = BytesIO(file_data)

        # --- PDF Handling ---
        if filename.endswith(".pdf"):
            result["file_type"] = "PDF"
            try:
                reader = PdfReader(buffer)
                result["num_pages"] = len(reader.pages)
                buffer.seek(0)
                images = extract_images(buffer)
                result["num_images"] = len(images)
                result["images"] = images
                result["is_pdf"] = True
            except Exception as e:
                result["errors"].append(f"Invalid PDF: {str(e)}")

        # --- DOCX Handling ---
        elif filename.endswith(".docx"):
            result["file_type"] = "DOCX"
            try:
                buffer.seek(0)
                doc = Document(buffer)
                # Approximate "pages" = 1 for every 800 words
                text = "\n".join([p.text for p in doc.paragraphs])
                word_count = len(text.split())
                result["num_pages"] = max(1, word_count // 800)
                # Count images
                result["num_images"] = len(doc.inline_shapes)
                result["is_docx"] = True
            except Exception as e:
                result["errors"].append(f"Invalid DOCX: {str(e)}")

        # --- Unsupported ---
        else:
            result["errors"].append("Unsupported file type")

    except Exception as e:
        result["errors"].append(f"File processing error: {str(e)}")

    return result


def extract_images(buffer):
    """Extract images inside a PDF (PyMuPDF)."""

    images = []
    try:
        # Ensure buffer is at start
        if hasattr(buffer, "seek"):
            buffer.seek(0)

        pdf_doc = fitz.open(stream=buffer.read(), filetype="pdf")
        for page_index, page in enumerate(pdf_doc, start=1):
            try:
                # PyMuPDF v1.23+ uses get_images(), older uses getImageList()
                image_list = page.get_images(full=True) if hasattr(page, "get_images") else page.getImageList()

                for img_index, img in enumerate(image_list, start=1):
                    xref = img[0]
                    base_image = pdf_doc.extract_image(xref)
                    image_data = base_image["image"]
                    images.append((BytesIO(image_data), page_index))

            except Exception as page_err:
                print(f"[WARN] Page {page_index}: Failed to extract some images: {page_err}")

    except Exception as e:
        print(f"[WARN] Failed to extract images: {e}")

    return images

#docx files do not have pages, so we can only approximate them with 
from io import BytesIO
from docx import Document

from io import BytesIO
from docx import Document

def extract_docx_content(file_obj):
    """
    Extract text and embedded images from a DOCX file.
    Approximates 1 page = 800 words.
    Returns:
        text (str): Combined document text.
        images (list[tuple[BytesIO, int]]): List of (image_data, approx_page_index) tuples.
    """
    # Ensure we're at the start
    if hasattr(file_obj, "seek"):
        file_obj.seek(0)

    # Load file into memory buffer
    if hasattr(file_obj, "read"):
        buffer = BytesIO(file_obj.read())
    else:
        buffer = open(file_obj, "rb")

    doc = Document(buffer)

    all_text = []
    images = []
    word_count = 0

    # namespace key used to read rId attribute
    REL_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"

    # iterate paragraphs in order to preserve document flow
    for para in doc.paragraphs:
        para_text = para.text or ""
        # approximate page for images that appear in this paragraph
        approx_page = max(1, word_count // 400 + 1)

        # search every run inside paragraph for image blip references
        for run in para.runs:
            # find any <blip ... r:embed="rIdXYZ" /> nodes under the run XML
            blips = run.element.findall('.//{*}blip')
            for blip in blips:
                rId = blip.get(REL_NS)
                if not rId:
                    continue
                # fetch image blob from document relationships if present
                if rId in doc.part.rels:
                    try:
                        image_blob = doc.part.rels[rId].target_part.blob
                        images.append((BytesIO(image_blob), approx_page))
                    except Exception:
                        # skip if extracting that particular image fails
                        continue

        # accumulate text and update word count after scanning runs (so images in this para map to current word_count)
        if para_text.strip():
            all_text.append(para_text)
            word_count += len(para_text.split())

    # cleanup if we opened a file object (not necessary for BytesIO)
    if not hasattr(file_obj, "read"):
        buffer.close()

    return "\n".join(all_text), images


